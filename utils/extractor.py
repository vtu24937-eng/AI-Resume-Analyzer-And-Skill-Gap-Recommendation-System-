from pypdf import PdfReader
import docx
import os


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pypdf (pure Python)."""
    text = ""
    try:
        reader = PdfReader(filepath)
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    except Exception as e:
        text = f"[Error reading PDF: {e}]"
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        document = docx.Document(filepath)
        for para in document.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        text = f"[Error reading DOCX: {e}]"
    return text.strip()


def extract_text(filepath: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        return "[Unsupported file format]"
